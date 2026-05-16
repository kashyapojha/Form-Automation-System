# =========================================================

# Excel-Driven Automated Form Population Engine

# =========================================================

# This script:

# 1. Reads data from Excel

# 2. Generates employee-specific Word forms automatically

# 3. Fills all dynamic fields

# =========================================================



import pandas as pd

from docx import Document

from docx.shared import Pt

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT



# =========================================================

# LOAD EXCEL SHEETS

# =========================================================



excel_file = "Form_Input_Template.xlsx"



employee_df = pd.read_excel(excel_file, sheet_name="Employee_Master")

device_df = pd.read_excel(excel_file, sheet_name="Device_Details")

issue_df = pd.read_excel(excel_file, sheet_name="Issue_Details")



# =========================================================

# GENERATE FORM FOR EACH EMPLOYEE

# =========================================================



for index, employee in employee_df.iterrows():



    employee_id = str(employee["Employee ID"])

    issue_rows = issue_df[issue_df["Employee ID"].astype(str) == employee_id]
    device_rows = device_df[device_df["Employee ID"].astype(str) == employee_id]

    if issue_rows.empty or device_rows.empty:
        missing = []
        if issue_rows.empty:
            missing.append("Issue_Details")
        if device_rows.empty:
            missing.append("Device_Details")
        print(f"Skipped {employee_id}: no row in {', '.join(missing)}. Add matching rows in Excel.")
        continue

    issue_info = issue_rows.iloc[0]
    devices = device_rows



    # =====================================================

    # CREATE WORD DOCUMENT

    # =====================================================



    doc = Document()



    # =====================================================

    # TITLE

    # =====================================================



    title = doc.add_paragraph()

    run = title.add_run("ABC")

    run.bold = True

    run.font.size = Pt(18)

    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



    subtitle = doc.add_paragraph()

    run = subtitle.add_run("IT Department — Device Issue & Return Form")

    run.bold = True

    run.font.size = Pt(14)

    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



    doc.add_paragraph(f"Form No.: {employee_id}")

    doc.add_paragraph("Date: 15-May-2025")



    # =====================================================

    # EMPLOYEE INFORMATION

    # =====================================================



    doc.add_heading("EMPLOYEE INFORMATION", level=1)



    emp_table = doc.add_table(rows=6, cols=2)

    emp_table.style = "Table Grid"



    employee_fields = [

        ("Employee Name", employee["Employee Name"]),

        ("Employee ID", employee["Employee ID"]),

        ("Department", employee["Department"]),

        ("Designation", employee["Designation"]),

        ("Contact No.", employee["Contact No."]),

        ("Email", employee["Email"]),

    ]



    for row, (field, value) in zip(emp_table.rows, employee_fields):

        row.cells[0].text = str(field)

        row.cells[1].text = str(value)



    # =====================================================

    # DEVICE DETAILS

    # =====================================================



    doc.add_heading("DEVICE DETAILS", level=1)



    device_table = doc.add_table(rows=1, cols=6)

    device_table.style = "Table Grid"



    headers = [

        "#",

        "Device Type",

        "Brand / Model",

        "Serial No.",

        "Condition",

        "Remarks"

    ]



    hdr_cells = device_table.rows[0].cells



    for i, header in enumerate(headers):

        hdr_cells[i].text = header



    for idx, (_, device) in enumerate(devices.iterrows(), start=1):



        row_cells = device_table.add_row().cells



        row_cells[0].text = str(idx)

        row_cells[1].text = str(device["Device Type"])

        row_cells[2].text = str(device["Brand / Model"])

        row_cells[3].text = str(device["Serial No."])

        row_cells[4].text = str(device["Condition"])

        row_cells[5].text = str(device["Remarks"])



    # =====================================================

    # ISSUE DETAILS

    # =====================================================



    doc.add_heading("ISSUE & RETURN DETAILS", level=1)



    issue_table = doc.add_table(rows=4, cols=2)

    issue_table.style = "Table Grid"



    issue_fields = [

        ("Issue Date", issue_info["Issue Date"]),

        ("Expected Return Date", issue_info["Expected Return Date"]),

        ("Purpose of Issue", issue_info["Purpose of Issue"]),

        ("Approved By", issue_info["Approved By"]),

    ]



    for row, (field, value) in zip(issue_table.rows, issue_fields):

        row.cells[0].text = str(field)

        row.cells[1].text = str(value)



    # =====================================================

    # TERMS & CONDITIONS

    # =====================================================



    doc.add_heading("TERMS & CONDITIONS", level=1)



    terms = [

        "The employee is responsible for the safe custody of issued devices.",

        "Any damage or loss must be reported to the IT department immediately.",

        "Devices must be returned on or before the expected return date.",

        "Unauthorized use or transfer of devices is strictly prohibited."

    ]



    for term in terms:

        doc.add_paragraph(term, style="List Number")



    # =====================================================

    # SIGNATURE SECTION

    # =====================================================



    doc.add_heading("SIGNATURES", level=1)



    sig_table = doc.add_table(rows=3, cols=3)

    sig_table.style = "Table Grid"



    headers = ["Employee", "IT Manager", "HOD Approval"]



    for i, header in enumerate(headers):

        sig_table.rows[0].cells[i].text = header



    for i in range(3):

        sig_table.rows[1].cells[i].text = "Name: ___________________"

        sig_table.rows[2].cells[i].text = "Date: ____________________"



    # =====================================================

    # FOOTER

    # =====================================================



    footer = doc.add_paragraph()



    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



    run = footer.add_run(

        "ABC — IT Department\n"

        "Confidential | For Internal Use Only\n"

        "This is a system-generated document."

    )



    run.italic = True

    run.font.size = Pt(10)



    # =====================================================

    # SAVE FILE

    # =====================================================



    output_filename = f"{employee_id}_Device_Form.docx"



    doc.save(output_filename)



    print(f"Generated: {output_filename}")



# =========================================================

# END

# =========================================================


