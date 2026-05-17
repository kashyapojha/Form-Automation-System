"""Excel validation and DOCX form generation."""

from __future__ import annotations

import io
import zipfile
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import BinaryIO

import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

REQUIRED_SHEETS = {
    "Employee_Master": [
        "Employee ID",
        "Employee Name",
        "Department",
        "Designation",
        "Contact No.",
        "Email",
    ],
    "Device_Details": [
        "Employee ID",
        "Device Type",
        "Brand / Model",
        "Serial No.",
        "Condition",
        "Remarks",
    ],
    "Issue_Details": [
        "Employee ID",
        "Issue Date",
        "Expected Return Date",
        "Purpose of Issue",
        "Approved By",
    ],
}

TERMS = [
    "The employee is responsible for the safe custody of issued devices.",
    "Any damage or loss must be reported to the IT department immediately.",
    "Devices must be returned on or before the expected return date.",
    "Unauthorized use or transfer of devices is strictly prohibited.",
]


@dataclass
class ValidationIssue:
    employee_id: str
    message: str


@dataclass
class ValidationResult:
    valid: bool
    total_employees: int = 0
    ready_count: int = 0
    issues: list[ValidationIssue] = field(default_factory=list)
    missing_sheets: list[str] = field(default_factory=list)
    missing_columns: dict[str, list[str]] = field(default_factory=dict)


@dataclass
class GenerationResult:
    generated: list[tuple[str, bytes]]
    skipped: list[ValidationIssue]
    total_employees: int


def load_excel(source: BinaryIO | bytes | str | Path) -> dict[str, pd.DataFrame]:
    if isinstance(source, bytes):
        source = io.BytesIO(source)
    return pd.read_excel(source, sheet_name=list(REQUIRED_SHEETS.keys()))


def _check_structure(sheets: dict[str, pd.DataFrame]) -> ValidationResult:
    result = ValidationResult(valid=True, total_employees=0)

    for sheet_name, required_cols in REQUIRED_SHEETS.items():
        if sheet_name not in sheets:
            result.valid = False
            result.missing_sheets.append(sheet_name)
            continue
        df = sheets[sheet_name]
        missing = [c for c in required_cols if c not in df.columns]
        if missing:
            result.valid = False
            result.missing_columns[sheet_name] = missing

    if not result.valid:
        return result

    employee_df = sheets["Employee_Master"]
    issue_df = sheets["Issue_Details"]
    device_df = sheets["Device_Details"]
    result.total_employees = len(employee_df)

    for _, employee in employee_df.iterrows():
        employee_id = str(employee["Employee ID"]).strip()
        issue_rows = issue_df[issue_df["Employee ID"].astype(str).str.strip() == employee_id]
        device_rows = device_df[device_df["Employee ID"].astype(str).str.strip() == employee_id]

        if issue_rows.empty or device_rows.empty:
            missing_parts = []
            if issue_rows.empty:
                missing_parts.append("Issue_Details")
            if device_rows.empty:
                missing_parts.append("Device_Details")
            result.issues.append(
                ValidationIssue(
                    employee_id=employee_id,
                    message=f"Missing row in {', '.join(missing_parts)}",
                )
            )
        else:
            result.ready_count += 1

    result.valid = result.valid and result.ready_count > 0
    return result


def validate_excel(source: BinaryIO | bytes | str | Path) -> ValidationResult:
    sheets = load_excel(source)
    return _check_structure(sheets)


def _format_cell(value) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, (datetime, pd.Timestamp)):
        return value.strftime("%d-%b-%Y")
    return str(value)


def build_form_document(
    employee: pd.Series,
    devices: pd.DataFrame,
    issue_info: pd.Series,
    company_name: str = "ABC",
) -> Document:
    employee_id = str(employee["Employee ID"]).strip()
    doc = Document()

    title = doc.add_paragraph()
    run = title.add_run(company_name)
    run.bold = True
    run.font.size = Pt(18)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    run = subtitle.add_run("IT Department — Device Issue & Return Form")
    run.bold = True
    run.font.size = Pt(14)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(f"Form No.: {employee_id}")
    doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%b-%Y')}")

    doc.add_heading("EMPLOYEE INFORMATION", level=1)
    emp_table = doc.add_table(rows=6, cols=2)
    emp_table.style = "Table Grid"
    employee_fields = [
        ("Employee Name", employee["Employee Name"]),
        ("Employee ID", employee["Employee ID"]),
        ("Department", employee["Department"]),
        ("Designation", employee["Designation"]),
        ("Contact No.", employee["Contact No."]),
        ("Email", employee["Email"]),
    ]
    for row, (field, value) in zip(emp_table.rows, employee_fields):
        row.cells[0].text = str(field)
        row.cells[1].text = _format_cell(value)

    doc.add_heading("DEVICE DETAILS", level=1)
    device_table = doc.add_table(rows=1, cols=6)
    device_table.style = "Table Grid"
    headers = ["#", "Device Type", "Brand / Model", "Serial No.", "Condition", "Remarks"]
    hdr_cells = device_table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    for idx, (_, device) in enumerate(devices.iterrows(), start=1):
        row_cells = device_table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = _format_cell(device["Device Type"])
        row_cells[2].text = _format_cell(device["Brand / Model"])
        row_cells[3].text = _format_cell(device["Serial No."])
        row_cells[4].text = _format_cell(device["Condition"])
        row_cells[5].text = _format_cell(device["Remarks"])

    doc.add_heading("ISSUE & RETURN DETAILS", level=1)
    issue_table = doc.add_table(rows=4, cols=2)
    issue_table.style = "Table Grid"
    issue_fields = [
        ("Issue Date", issue_info["Issue Date"]),
        ("Expected Return Date", issue_info["Expected Return Date"]),
        ("Purpose of Issue", issue_info["Purpose of Issue"]),
        ("Approved By", issue_info["Approved By"]),
    ]
    for row, (field, value) in zip(issue_table.rows, issue_fields):
        row.cells[0].text = str(field)
        row.cells[1].text = _format_cell(value)

    doc.add_heading("TERMS & CONDITIONS", level=1)
    for term in TERMS:
        doc.add_paragraph(term, style="List Number")

    doc.add_heading("SIGNATURES", level=1)
    sig_table = doc.add_table(rows=3, cols=3)
    sig_table.style = "Table Grid"
    sig_headers = ["Employee", "IT Manager", "HOD Approval"]
    for i, header in enumerate(sig_headers):
        sig_table.rows[0].cells[i].text = header
    for i in range(3):
        sig_table.rows[1].cells[i].text = "Name: ___________________"
        sig_table.rows[2].cells[i].text = "Date: ____________________"

    footer = doc.add_paragraph()
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = footer.add_run(
        f"{company_name} — IT Department\n"
        "Confidential | For Internal Use Only\n"
        "This is a system-generated document."
    )
    run.italic = True
    run.font.size = Pt(10)

    return doc


def generate_forms(
    source: BinaryIO | bytes | str | Path,
    company_name: str = "ABC",
) -> GenerationResult:
    sheets = load_excel(source)
    employee_df = sheets["Employee_Master"]
    issue_df = sheets["Issue_Details"]
    device_df = sheets["Device_Details"]

    generated: list[tuple[str, bytes]] = []
    skipped: list[ValidationIssue] = []

    for _, employee in employee_df.iterrows():
        employee_id = str(employee["Employee ID"]).strip()
        issue_rows = issue_df[issue_df["Employee ID"].astype(str).str.strip() == employee_id]
        device_rows = device_df[device_df["Employee ID"].astype(str).str.strip() == employee_id]

        if issue_rows.empty or device_rows.empty:
            missing_parts = []
            if issue_rows.empty:
                missing_parts.append("Issue_Details")
            if device_rows.empty:
                missing_parts.append("Device_Details")
            skipped.append(
                ValidationIssue(
                    employee_id=employee_id,
                    message=f"Missing row in {', '.join(missing_parts)}",
                )
            )
            continue

        doc = build_form_document(employee, device_rows, issue_rows.iloc[0], company_name)
        buffer = io.BytesIO()
        doc.save(buffer)
        safe_id = "".join(c if c.isalnum() or c in "-_" else "_" for c in employee_id)
        filename = f"{safe_id}_Device_Form.docx"
        generated.append((filename, buffer.getvalue()))

    return GenerationResult(
        generated=generated,
        skipped=skipped,
        total_employees=len(employee_df),
    )


def create_zip_archive(files: list[tuple[str, bytes]]) -> bytes:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for filename, data in files:
            zf.writestr(filename, data)
    return buffer.getvalue()
